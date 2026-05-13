"""
tools/file_tools.py
===================

File-reading helpers for resumes and job descriptions.

Supports:
  • .txt / .md  — read verbatim
  • .pdf        — `pypdf`
  • .docx       — `python-docx`

Kept dependency-light: any failure in the optional extractors degrades to
returning an empty string so the agent can keep running even if one resume
is malformed.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Union

from config import get_logger

log = get_logger(__name__)

PathLike = Union[str, Path]


# ---------------------------------------------------------------------------
# Public helpers
# ---------------------------------------------------------------------------
def read_resume_text(path: PathLike) -> str:
    """Return the plain-text content of a resume file."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    raise ValueError(f"Unsupported resume format: {suffix}")


def read_jd_text(path: PathLike) -> str:
    """Return the plain-text content of a JD file (same backends as resumes)."""
    return read_resume_text(path)


def read_uploaded_file(name: str, data: bytes) -> str:
    """
    Read an uploaded file's bytes into plain text.

    Used by the Streamlit UI where files arrive as in-memory buffers.
    """
    suffix = Path(name).suffix.lower()
    if suffix in {".txt", ".md"}:
        return data.decode("utf-8", errors="ignore")
    if suffix == ".pdf":
        return _read_pdf_bytes(data)
    if suffix == ".docx":
        return _read_docx_bytes(data)
    # Fall back to UTF-8 best-effort
    return data.decode("utf-8", errors="ignore")


# ---------------------------------------------------------------------------
# Backend-specific readers
# ---------------------------------------------------------------------------
def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # noqa: BLE001
        log.warning("pypdf not available: %s", exc)
        return ""
    try:
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # noqa: BLE001
        log.warning("Failed to read PDF %s: %s", path, exc)
        return ""


def _read_pdf_bytes(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception:  # noqa: BLE001
        return ""
    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # noqa: BLE001
        log.warning("Failed to parse uploaded PDF: %s", exc)
        return ""


def _read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except Exception as exc:  # noqa: BLE001
        log.warning("python-docx not available: %s", exc)
        return ""
    try:
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as exc:  # noqa: BLE001
        log.warning("Failed to read DOCX %s: %s", path, exc)
        return ""


def _read_docx_bytes(data: bytes) -> str:
    try:
        import docx  # python-docx
    except Exception:  # noqa: BLE001
        return ""
    try:
        doc = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as exc:  # noqa: BLE001
        log.warning("Failed to parse uploaded DOCX: %s", exc)
        return ""


# ---------------------------------------------------------------------------
# Small write helpers (used by the Streamlit upload UI)
# ---------------------------------------------------------------------------
def save_resume_text(directory: PathLike, name: str, text: str) -> Path:
    """Persist a resume as plain text in the resumes directory."""
    directory = Path(directory)
    directory.mkdir(parents=True, exist_ok=True)
    safe = name.replace("/", "_").replace("\\", "_")
    if not Path(safe).suffix:
        safe += ".txt"
    out = directory / safe
    out.write_text(text, encoding="utf-8")
    log.info("Saved resume → %s", out)
    return out
